"""DOCX rendering for a Career Vault master resume.

This renderer intentionally uses only approved, already-generated draft content. It does
not make model calls and never reads data belonging to another user.
"""

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(33, 101, 69)
INK = RGBColor(24, 34, 31)
MUTED = RGBColor(81, 96, 88)


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_run_font(run: Any, size: float, bold: bool = False, color: RGBColor = INK) -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def _add_section_heading(document: Document, text: str, size: float = 9) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    _set_run_font(run, size, bold=True, color=ACCENT)


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        if not isinstance(item, str) or not item.strip():
            continue
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(item.strip())
        _set_run_font(run, 9.5)


def _clean_list(values: Any) -> list[str]:
    return [str(value).strip() for value in values if str(value).strip()] if isinstance(values, list) else []


def render_master_resume(path: Path, display_name: str, resume: dict[str, Any]) -> None:
    """Create a compact, ATS-readable DOCX with a resume contact header override."""
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = INK

    # Named resume-contact-header override: the identity is prominent, while the body
    # stays entirely text-based and easy for ATS systems to read.
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    run = name.add_run(display_name.upper())
    _set_run_font(run, 20, bold=True, color=ACCENT)
    headline = str(resume.get("headline") or "Professional Profile").strip()
    if headline:
        line = document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(10)
        run = line.add_run(headline)
        _set_run_font(run, 10.5, color=MUTED)

    summary = str(resume.get("professional_summary") or "").strip()
    if summary:
        _add_section_heading(document, "Professional Summary")
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.08
        run = paragraph.add_run(summary)
        _set_run_font(run, 9.5)

    skills = _clean_list(resume.get("skills"))
    if skills:
        _add_section_heading(document, "Core Skills")
        table = document.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        _set_cell_shading(cell, "F0F6F1")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.space_before = Pt(2)
        run = paragraph.add_run("  •  ".join(skills))
        _set_run_font(run, 9.3, color=MUTED)

    experience = resume.get("experience")
    if isinstance(experience, list) and experience:
        _add_section_heading(document, "Professional Experience")
        for entry in experience:
            if not isinstance(entry, dict):
                continue
            role = str(entry.get("role") or "Professional experience").strip()
            company = str(entry.get("company") or "").strip()
            dates = str(entry.get("dates") or "").strip()
            title = document.add_paragraph()
            title.paragraph_format.space_before = Pt(4)
            title.paragraph_format.space_after = Pt(1)
            title.paragraph_format.keep_with_next = True
            run = title.add_run(role)
            _set_run_font(run, 10, bold=True)
            if company:
                run = title.add_run(f" | {company}")
                _set_run_font(run, 10, color=MUTED)
            if dates:
                run = title.add_run(f"  {dates}")
                _set_run_font(run, 9, color=MUTED)
            _add_bullets(document, _clean_list(entry.get("highlights")))

    for section_key, heading in (("certifications", "Certifications"), ("education", "Education"), ("additional_highlights", "Additional Highlights")):
        values = _clean_list(resume.get(section_key))
        if values:
            _add_section_heading(document, heading)
            _add_bullets(document, values)

    document.save(path)


def _add_bottom_border(paragraph: Any, color: str, size: int = 6) -> None:
    """Add a thin bottom border under a paragraph, used as a section/header divider."""
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    properties.append(borders)


def _add_split_line(document: Document, left_text: str, left_size: float, left_bold: bool, left_color: RGBColor, right_text: str, right_size: float, right_color: RGBColor, space_before: float, space_after: float) -> None:
    """One line with the role/company on the left and dates flush right, via a right tab stop."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.keep_with_next = True
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
    run = paragraph.add_run(left_text)
    _set_run_font(run, left_size, bold=left_bold, color=left_color)
    if right_text:
        run = paragraph.add_run(f"\t{right_text}")
        _set_run_font(run, right_size, color=right_color)


def _resume_density(resume: dict[str, Any]) -> dict[str, float]:
    """Pick tighter fonts/spacing as content grows, so a tailored resume stays a realistic
    single page instead of overflowing. This is an approximation, not a guarantee — the
    resume content itself is already capped upstream to a one-page-worthy amount."""
    experience = [entry for entry in resume.get("experience", []) if isinstance(entry, dict)]
    bullet_count = sum(len(_clean_list(entry.get("highlights"))) for entry in experience)
    extra_sections = len(_clean_list(resume.get("certifications"))) + len(_clean_list(resume.get("education"))) + len(_clean_list(resume.get("additional_highlights")))
    weight = bullet_count + len(experience) * 1.5 + extra_sections * 0.4

    if weight <= 11:
        return {"margin": 0.62, "name": 21, "headline": 10.8, "heading": 9.3, "body": 10, "bullet": 9.8, "bullet_gap": 2.4, "section_gap": 12}
    if weight <= 16:
        return {"margin": 0.55, "name": 19.5, "headline": 10.2, "heading": 9, "body": 9.4, "bullet": 9.2, "bullet_gap": 1.6, "section_gap": 9}
    return {"margin": 0.48, "name": 18, "headline": 9.7, "heading": 8.6, "body": 8.9, "bullet": 8.7, "bullet_gap": 1, "section_gap": 7}


def render_tailored_resume(path: Path, display_name: str, resume: dict[str, Any]) -> None:
    """Create a polished, industry-standard, single-page-oriented DOCX for a job-tailored resume.

    Unlike render_master_resume (which favors completeness), this renderer assumes the resume
    content has already been curated to a small, job-relevant set of entries and adapts font
    size/spacing to that content so the result reads like a real one-page resume: name header
    with a divider, right-aligned dates on each role line, and tight, scannable bullets.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    density = _resume_density(resume)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(density["margin"])
    section.bottom_margin = Inches(density["margin"])
    section.left_margin = Inches(density["margin"])
    section.right_margin = Inches(density["margin"])

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(density["body"])
    styles["Normal"].font.color.rgb = INK

    name = document.add_paragraph()
    name.paragraph_format.space_after = Pt(1)
    run = name.add_run(display_name.strip().upper() or "CANDIDATE")
    _set_run_font(run, density["name"], bold=True, color=ACCENT)

    headline = str(resume.get("headline") or "").strip()
    divider = document.add_paragraph()
    divider.paragraph_format.space_after = Pt(density["section_gap"] * 0.7)
    if headline:
        run = divider.add_run(headline)
        _set_run_font(run, density["headline"], color=MUTED)
    _add_bottom_border(divider, "CDDDCF", size=10)

    summary = str(resume.get("professional_summary") or "").strip()
    if summary:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(density["section_gap"])
        paragraph.paragraph_format.line_spacing = 1.12
        run = paragraph.add_run(summary)
        _set_run_font(run, density["body"])

    skills = _clean_list(resume.get("skills"))
    if skills:
        _add_section_heading(document, "Core Skills", density["heading"])
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(density["section_gap"])
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run("  •  ".join(skills))
        _set_run_font(run, density["bullet"], color=MUTED)

    experience = resume.get("experience")
    if isinstance(experience, list) and experience:
        _add_section_heading(document, "Professional Experience", density["heading"])
        for entry in experience:
            if not isinstance(entry, dict):
                continue
            role = str(entry.get("role") or "Professional experience").strip()
            company = str(entry.get("company") or "").strip()
            dates = str(entry.get("dates") or "").strip()
            left_text = f"{role} | {company}" if company else role
            _add_split_line(document, left_text, density["body"], True, INK, dates, density["bullet"], MUTED, 5, 1)
            for item in _clean_list(entry.get("highlights")):
                bullet = document.add_paragraph(style="List Bullet")
                bullet.paragraph_format.space_after = Pt(density["bullet_gap"])
                bullet.paragraph_format.line_spacing = 1.05
                run = bullet.add_run(item)
                _set_run_font(run, density["bullet"])
        document.paragraphs[-1].paragraph_format.space_after = Pt(density["section_gap"])

    for section_key, heading in (("certifications", "Certifications"), ("education", "Education"), ("additional_highlights", "Additional Highlights")):
        values = _clean_list(resume.get(section_key))
        if values:
            _add_section_heading(document, heading, density["heading"])
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(density["section_gap"])
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run("  •  ".join(values))
            _set_run_font(run, density["bullet"], color=MUTED)

    document.save(path)


def render_cover_letter(path: Path, display_name: str, job_title: str, cover_letter_body: str) -> None:
    """Create a simple, ATS-readable DOCX cover letter matching the resume's styling."""
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = INK

    name = document.add_paragraph()
    name.paragraph_format.space_after = Pt(2)
    run = name.add_run(display_name.strip() or "Candidate")
    _set_run_font(run, 15, bold=True, color=ACCENT)

    if job_title.strip():
        role_line = document.add_paragraph()
        role_line.paragraph_format.space_after = Pt(20)
        run = role_line.add_run(f"Application for: {job_title.strip()}")
        _set_run_font(run, 10, color=MUTED)

    for paragraph_text in [line.strip() for line in cover_letter_body.split("\n") if line.strip()]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.2
        run = paragraph.add_run(paragraph_text)
        _set_run_font(run, 10.5)

    document.save(path)
